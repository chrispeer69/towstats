"""morning_report -- the 6 AM printable day-prior report.

WHY THIS EXISTS SEPARATELY FROM notifier.py
    notifier.py renders HTML for email clients. This renders a .docx the owner
    opens from a folder and prints. Different medium, different constraints:
    email is scrollable and disposable, a printed page is finite and gets read
    once. The hard limit here is THREE PAGES -- a report nobody finishes is a
    report that was never written.

SCOPE IS ONE DAY. The subject of the report is the previous calendar day and
nothing else. Prior days appear only as a baseline strip beside each headline
number, because "41 wins" is not information until you know whether 41 is good.

TWO CORRECTNESS RULES THAT ARE EASY TO GET WRONG
    1. offered_at is naive UTC in the database (ingestion.py: "Everything is
       converted to UTC on the way into the database"). Every hour shown to a
       human MUST be converted to local time first. Reading the stored hour
       directly puts the peak at 13:00 when it is really 09:00 -- a four-hour
       error that would staff the wrong shift.
    2. "Yesterday" is a LOCAL midnight-to-midnight window, converted to UTC for
       the query. Filtering on date(offered_at) silently reports a day that
       starts at 8 PM.

NO DOLLAR FIGURES. offerAmount is empty on 100% of the records this account
produces (verified: sum(amount) == 0 over 3,693 rows). Every count is jobs, and
no total may be phrased so a reader infers revenue.
"""

from __future__ import annotations

import sqlite3
from collections import Counter
from dataclasses import dataclass, field
from datetime import date, datetime, time, timedelta, timezone
from pathlib import Path
from zoneinfo import ZoneInfo

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# -- outcome vocabulary ----------------------------------------------------
#
# WIN is an accepted offer. EVERY OTHER SETTLED OFFER IS A LOSS, cancellations
# included. A cancelled job is one we did not do, and the owner's rule is that
# the report counts it that way.
#
# This is a deliberate departure from MISSED_WORK_MODEL.md, which keeps client
# withdrawals out of the recoverable inventory. The evidence says the exclusion
# was hiding real failures rather than protecting a fair number:
#
#   * Over the record, excluding cancellations reports a 50.1% win rate. The
#     true share of offers worked is 38.0% -- a 12-point flattery.
#   * Allstate is the extreme case: 515 of its 965 offers cancel. Excluding
#     them printed Allstate at 74% (and at 100% on individual days, with zero
#     losses) while barely a third of its work was actually run.
#   * The cancellations are not all the client changing its mind. Every one of
#     the 39 "Service Failure Confirmed" had a driver assigned -- we took the
#     job and did not deliver it. 52 "Rejected By Motor Club" carry OUR stated
#     reason, mostly "No Drivers Available". Those are capacity failures.
#
# So cancellations count as losses, and the report then says WHY each one
# happened -- see CANCEL_KINDS. A total that hides its causes would just move
# the problem.
WIN_STATUSES = ("accepted",)
LOSS_STATUSES = ("denied", "expired", "canceled")

#: How Towbook's raw cancellation labels map onto causes a human can act on.
#: Ordered most-actionable first; that is the order the report prints them in.
#: ``blame`` marks the ones inside our control, which the report calls out.
CANCEL_KINDS: tuple[tuple[str, str, bool], ...] = (
    ("Service Failure Confirmed", "We failed after accepting", True),
    ("Rejected By Motor Club", "Club pulled it back", True),
    ("Goa Approved By Motor Club", "Gone on arrival", False),
    ("Service No Longer Needed", "Customer stood down", False),
    ("Cancelled", "Cancelled, no reason recorded", False),
)
CANCEL_OTHER = "Other / unrecognised"

#: An hour needs at least this many offers before it may be called strongest or
#: weakest. A single day spreads ~100 offers over 24 hours, so a 1-offer hour
#: sits at 0% or 100% and would win every superlative on noise alone.
MIN_HOURS_VOLUME = 3

#: Colours. Muted enough to print legibly on a mono laser printer.
INK = RGBColor(0x1C, 0x1C, 0x1C)
MUTED = RGBColor(0x6A, 0x6A, 0x6A)
WIN_C = RGBColor(0x1A, 0x7F, 0x37)
LOSS_C = RGBColor(0xB4, 0x23, 0x18)
RULE = "D9D9D9"


@dataclass
class HourRow:
    """One local hour of the reported day."""

    hour: int
    offers: int = 0
    wins: int = 0
    losses: int = 0
    canceled: int = 0

    @property
    def decided(self) -> int:
        return self.wins + self.losses

    @property
    def rate(self) -> float | None:
        return (self.wins / self.decided) if self.decided else None

    @property
    def label(self) -> str:
        suffix = "a" if self.hour < 12 else "p"
        base = self.hour % 12 or 12
        return f"{base}{suffix}"


@dataclass
class DayStats:
    """Everything the report says about one calendar day."""

    day: date
    offers: int = 0
    wins: int = 0
    losses: int = 0
    canceled: int = 0
    no_answer: int = 0
    declined: int = 0
    decline_reasons: Counter = field(default_factory=Counter)
    #: cancellation cause -> count, and the subset that already had a driver
    #: assigned. A cancellation we had crewed is a truck-roll or a commitment
    #: broken; one we never touched is a different problem with a different fix.
    cancel_kinds: Counter = field(default_factory=Counter)
    cancel_engaged: Counter = field(default_factory=Counter)
    cancel_reasons: Counter = field(default_factory=Counter)
    hours: list[HourRow] = field(default_factory=list)
    clients: dict[str, HourRow] = field(default_factory=dict)
    last_offer_utc: datetime | None = None

    @property
    def decided(self) -> int:
        return self.wins + self.losses

    @property
    def rate(self) -> float | None:
        return (self.wins / self.decided) if self.decided else None

    @property
    def has_data(self) -> bool:
        return self.offers > 0


# -- data access -----------------------------------------------------------


def _local_day_bounds(day: date, tz: ZoneInfo) -> tuple[datetime, datetime]:
    """Return the naive-UTC half-open [start, end) covering a LOCAL day.

    Built from local midnight so the window survives DST: on the spring-forward
    day this is 23 hours and on the fall-back day 25, which is exactly what
    "that day, locally" means.
    """
    start_local = datetime.combine(day, time.min, tzinfo=tz)
    end_local = datetime.combine(day + timedelta(days=1), time.min, tzinfo=tz)
    to_utc = lambda d: d.astimezone(timezone.utc).replace(tzinfo=None)  # noqa: E731
    return to_utc(start_local), to_utc(end_local)


def load_day(conn: sqlite3.Connection, day: date, tz: ZoneInfo) -> DayStats:
    """Compute one local day's statistics straight from ``requests``.

    Reads the raw offer rows rather than metrics_daily because the shape of
    that JSON blob is owned by metrics.py and this report needs a local-hour
    breakdown that the blob does not carry.
    """
    start, end = _local_day_bounds(day, tz)
    stats = DayStats(day=day, hours=[HourRow(h) for h in range(24)])

    rows = conn.execute(
        """
        SELECT offered_at, status, denial_reason, client_name, status_raw, driver_assigned
          FROM requests
         WHERE offered_at >= ? AND offered_at < ?
        """,
        (start.isoformat(sep=" "), end.isoformat(sep=" ")),
    ).fetchall()

    for offered_at, status, denial_reason, client_name, status_raw, driver in rows:
        stamp = _parse_utc(offered_at)
        if stamp is None:
            continue
        local = stamp.replace(tzinfo=timezone.utc).astimezone(tz)
        bucket = stats.hours[local.hour]
        client = stats.clients.setdefault(client_name or "Unknown", HourRow(-1))

        stats.offers += 1
        bucket.offers += 1
        client.offers += 1
        if stats.last_offer_utc is None or stamp > stats.last_offer_utc:
            stats.last_offer_utc = stamp

        if status in WIN_STATUSES:
            stats.wins += 1
            bucket.wins += 1
            client.wins += 1
        elif status in LOSS_STATUSES:
            stats.losses += 1
            bucket.losses += 1
            client.losses += 1
            if status == "expired":
                stats.no_answer += 1
            elif status == "denied":
                stats.declined += 1
                stats.decline_reasons[(denial_reason or "No reason given").strip()] += 1
            else:
                stats.canceled += 1
                bucket.canceled += 1
                client.canceled += 1
                kind = _cancel_kind(status_raw)
                stats.cancel_kinds[kind] += 1
                if driver:
                    stats.cancel_engaged[kind] += 1
                if denial_reason:
                    stats.cancel_reasons[denial_reason.strip()] += 1

    return stats


def _cancel_kind(status_raw: object) -> str:
    """Map a raw Towbook cancellation label onto an actionable cause.

    Matched case-insensitively, and anything unrecognised falls into a named
    bucket rather than being dropped -- if Towbook adds a label, it has to show
    up on the page as unclassified work, not silently vanish from the totals.
    """
    text = (status_raw or "").strip().lower()
    for raw, label, _blame in CANCEL_KINDS:
        if text == raw.lower():
            return label
    return CANCEL_OTHER


def _parse_utc(value: object) -> datetime | None:
    if isinstance(value, datetime):
        return value.replace(tzinfo=None) if value.tzinfo else value
    if not isinstance(value, str) or not value.strip():
        return None
    try:
        return datetime.fromisoformat(value.strip().replace("Z", "")).replace(tzinfo=None)
    except ValueError:
        return None


def same_weekday_baseline(
    conn: sqlite3.Connection, day: date, tz: ZoneInfo, count: int = 4
) -> tuple[list[DayStats], str]:
    """The ``count`` most recent same-weekday days that actually had offers.

    Same-weekday rather than trailing-N because tow volume is strongly weekly:
    comparing a Sunday to a Thursday manufactures a swing that is just the
    calendar. Days with no offers are skipped rather than averaged in as zeros,
    which would drag every baseline toward nothing during a gap in the feed.
    """
    found: list[DayStats] = []
    cursor = day - timedelta(days=7)
    # Bounded walk: look back at most 12 same-weekdays (~3 months) for `count`.
    for _ in range(12):
        if len(found) >= count:
            break
        stats = load_day(conn, cursor, tz)
        if stats.has_data:
            found.append(stats)
        cursor -= timedelta(days=7)
    label = f"{len(found)}-{day.strftime('%a')} avg" if found else "no baseline"
    return found, label


def _avg(values: list[float | None]) -> float | None:
    real = [v for v in values if v is not None]
    return sum(real) / len(real) if real else None


# -- docx helpers ----------------------------------------------------------
#
# python-docx exposes no API for cell shading or single-sided borders, so the
# two helpers below drop to the underlying XML. Both are deliberately tiny and
# used everywhere, so the raw-XML surface stays in one place.


def _shade(cell, hex_fill: str) -> None:
    from docx.oxml import OxmlElement

    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(el)


def _border(cell, edge: str, hex_color: str, size: int = 4) -> None:
    from docx.oxml import OxmlElement

    pr = cell._tc.get_or_add_tcPr()
    borders = pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        pr.append(borders)
    el = OxmlElement(f"w:{edge}")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), str(size))
    el.set(qn("w:color"), hex_color)
    borders.append(el)


def _text(
    cell_or_para,
    value: str,
    *,
    size: int = 9,
    bold: bool = False,
    color: RGBColor = INK,
    align: str | None = None,
    space_after: int = 0,
) -> None:
    """Write one run of text, replacing whatever the paragraph held."""
    para = cell_or_para.paragraphs[0] if hasattr(cell_or_para, "paragraphs") else cell_or_para
    para.text = ""
    run = para.add_run(value)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = "Segoe UI"
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.space_before = Pt(0)
    if align == "r":
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == "c":
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _pct(value: float | None) -> str:
    return "--" if value is None else f"{round(value * 100)}%"


def _delta(
    now: float | None,
    before: float | None,
    *,
    as_pct: bool = False,
    higher_is_better: bool = True,
) -> tuple[str, RGBColor]:
    """A signed comparison marker, or an em dash when there is no baseline.

    The arrow always points the way the raw number moved, but the COLOUR
    follows whether that move was good news. Losses are the case that matters:
    thirteen more of them is an arrow up and it must never print green, or the
    one number on the page you least want rising is the one that looks best.
    """
    if now is None or before is None:
        return "—", MUTED
    diff = now - before
    if abs(diff) < (0.005 if as_pct else 0.5):
        return "level", MUTED
    shown = f"{abs(round(diff * 100))} pts" if as_pct else (
        f"{abs(diff):.0f}" if abs(diff) >= 1 else f"{abs(diff):.1f}"
    )
    arrow = "▲" if diff > 0 else "▼"
    good = (diff > 0) if higher_is_better else (diff < 0)
    return f"{arrow} {shown}", (WIN_C if good else LOSS_C)


# -- narrative -------------------------------------------------------------


def _hour_findings(stats: DayStats) -> tuple[list[HourRow], list[HourRow], list[HourRow]]:
    """(strongest, weakest, busy-but-losing) hours, each already volume-gated.

    The third list is the one that pays for the report. An hour can have a poor
    win rate because it is dead, which costs nothing, or because it is busy and
    nobody is answering, which costs a job every time. Only the second kind is
    worth a staffing change, so it is separated out rather than buried in the
    weakest list.
    """
    eligible = [h for h in stats.hours if h.decided >= MIN_HOURS_VOLUME]
    if not eligible:
        return [], [], []
    ranked = sorted(eligible, key=lambda h: (h.rate or 0, h.offers))
    strongest = list(reversed(ranked[-3:]))
    weakest = ranked[:3]

    median_offers = sorted(h.offers for h in eligible)[len(eligible) // 2]
    busy_losing = sorted(
        (h for h in eligible if h.offers >= max(median_offers, MIN_HOURS_VOLUME) and (h.rate or 0) < 0.35),
        key=lambda h: (-(h.losses), -(h.offers)),
    )[:3]
    return strongest, weakest, busy_losing


def _headline(stats: DayStats, baseline: list[DayStats], busy_losing: list[HourRow]) -> str:
    """One plain sentence a person can act on, or repeat to a dispatcher."""
    if not stats.has_data:
        return "No offers were recorded for this day."

    base_rate = _avg([b.rate for b in baseline])
    parts: list[str] = []

    if base_rate is not None and stats.rate is not None:
        swing = round((stats.rate - base_rate) * 100)
        if abs(swing) >= 5:
            direction = "above" if swing > 0 else "below"
            parts.append(
                f"Win rate of {_pct(stats.rate)} ran {abs(swing)} points {direction} "
                f"the usual {stats.day.strftime('%A')} ({_pct(base_rate)})."
            )
        else:
            parts.append(f"Win rate of {_pct(stats.rate)} was normal for a {stats.day.strftime('%A')}.")
    else:
        parts.append(f"Won {stats.wins} of {stats.decided} decided offers ({_pct(stats.rate)}).")

    if busy_losing:
        worst = busy_losing[0]
        parts.append(
            f"The costliest stretch was {worst.label}: {worst.offers} offers came in and "
            f"{worst.losses} were lost."
        )
    # Name the single biggest bucket of losses, whichever it is. Ordered by
    # size rather than by a fixed preference so the sentence follows the day
    # instead of always telling the same story.
    if stats.losses:
        biggest, count = max(
            (("offers nobody answered", stats.no_answer),
             ("offers we turned down", stats.declined),
             ("cancellations", stats.canceled)),
            key=lambda pair: pair[1],
        )
        if count and count / stats.losses >= 0.4:
            parts.append(f"{count} of {stats.losses} losses were {biggest}.")

    ours = sum(stats.cancel_kinds[label] for _raw, label, blame in CANCEL_KINDS if blame)
    if ours:
        parts.append(f"{ours} of the cancellations were inside our control.")
    return " ".join(parts)


# -- rendering -------------------------------------------------------------


def _scoreboard(doc, stats: DayStats, prior: DayStats, baseline: list[DayStats]) -> None:
    """The four headline numbers, each with its thin baseline strip."""
    base_rate = _avg([b.rate for b in baseline])
    base_wins = _avg([float(b.wins) for b in baseline])
    weekday = stats.day.strftime("%a")
    n = len(baseline)
    avg_label = f"{n}-{weekday} avg" if n else "no baseline"

    cells = [
        ("OFFERS", str(stats.offers), _delta(float(stats.offers), float(prior.offers) if prior.has_data else None), f"{prior.offers} prior day" if prior.has_data else "no prior day", INK),
        ("WINS", str(stats.wins), _delta(float(stats.wins), float(prior.wins) if prior.has_data else None), f"{avg_label}: {base_wins:.0f}" if base_wins is not None else avg_label, WIN_C),
        ("LOSSES", str(stats.losses), _delta(float(stats.losses), float(prior.losses) if prior.has_data else None, higher_is_better=False), f"{stats.no_answer} no answer · {stats.declined} declined · {stats.canceled} cancelled", LOSS_C),
        ("WIN RATE", _pct(stats.rate), _delta(stats.rate, prior.rate, as_pct=True), f"{avg_label}: {_pct(base_rate)}" if base_rate is not None else avg_label, INK),
    ]

    table = doc.add_table(rows=4, cols=4)
    table.autofit = False
    for idx, (label, value, (marker, marker_c), foot, value_c) in enumerate(cells):
        col = table.columns[idx]
        col.width = Inches(1.72)
        for r in range(4):
            table.cell(r, idx).width = Inches(1.72)
        _text(table.cell(0, idx), label, size=8, bold=True, color=MUTED)
        _text(table.cell(1, idx), value, size=26, bold=True, color=value_c)
        _text(table.cell(2, idx), marker, size=9, bold=True, color=marker_c)
        _text(table.cell(3, idx), foot, size=8, color=MUTED)
        _border(table.cell(0, idx), "top", "1C1C1C", size=8)
        for r in range(4):
            table.cell(r, idx).vertical_alignment = WD_ALIGN_VERTICAL.TOP
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _section(doc, title: str, note: str | None = None) -> None:
    para = doc.add_paragraph()
    _text(para, title, size=11, bold=True, space_after=1)
    if note:
        sub = doc.add_paragraph()
        _text(sub, note, size=8, color=MUTED, space_after=3)


def _hourly_table(doc, stats: DayStats) -> None:
    """24 local hours as two side-by-side 12-row panels.

    Two panels rather than one 24-row table purely for page budget: stacked, the
    same data costs most of a page and pushes the trend section onto a fourth.
    """
    # 13 columns: a 6-column panel, a narrow spacer, then a second 6-column
    # panel.
    #
    # Lost is shown explicitly, and that is not padding. Rate is wins over
    # DECIDED offers, so an hour with 9 offers, 6 wins and 3 client
    # cancellations is a true 100%. Printed as "9 offers, 6 won, 100%" that
    # looks like broken arithmetic and the reader stops trusting the page.
    # Printed as "9 offers, 6 won, 0 lost, 100%" it reconciles on sight.
    peak = max((h.offers for h in stats.hours), default=0) or 1
    left_cols, spacer_col, right_cols = (0, 1, 2, 3, 4, 5), 6, (7, 8, 9, 10, 11, 12)
    panel_w = [0.38, 0.68, 0.34, 0.36, 0.36, 0.44]
    widths = panel_w + [0.16] + panel_w
    heads = ["Hr", "Volume", "Off", "Won", "Lost", "Rate"]
    numeric = {"Off", "Won", "Lost", "Rate"}

    table = doc.add_table(rows=13, cols=13)
    table.autofit = False
    for c, w in enumerate(widths):
        for r in range(13):
            table.cell(r, c).width = Inches(w)
    for panel in (left_cols, right_cols):
        for c, head in zip(panel, heads):
            _text(table.cell(0, c), head, size=8, bold=True, color=MUTED,
                  align="r" if head in numeric else None)
            _border(table.cell(0, c), "bottom", "1C1C1C")
    _text(table.cell(0, spacer_col), "", size=8)

    for row in range(12):
        for panel, hr in ((left_cols, stats.hours[row]), (right_cols, stats.hours[row + 12])):
            c_hr, c_bar, c_off, c_won, c_lost, c_rate = panel
            bar = "█" * round((hr.offers / peak) * 7)
            dead = hr.offers == 0
            cold = hr.rate is not None and hr.rate < 0.25 and hr.decided >= MIN_HOURS_VOLUME
            _text(table.cell(row + 1, c_hr), hr.label, size=8, color=MUTED)
            _text(table.cell(row + 1, c_bar), bar or "·", size=8, color=MUTED if dead else INK)
            _text(table.cell(row + 1, c_off), "·" if dead else str(hr.offers), size=8, align="r")
            _text(table.cell(row + 1, c_won), "·" if dead else str(hr.wins), size=8, align="r",
                  bold=bool(hr.wins), color=WIN_C if hr.wins else MUTED)
            _text(table.cell(row + 1, c_lost), "·" if dead else str(hr.losses), size=8, align="r",
                  color=LOSS_C if hr.losses else MUTED)
            _text(table.cell(row + 1, c_rate), _pct(hr.rate) if hr.decided else "·", size=8, align="r",
                  bold=cold, color=LOSS_C if cold else INK)
            _border(table.cell(row + 1, c_hr), "top", RULE)


def _bullets(doc, rows: list[tuple[str, str, RGBColor]]) -> None:
    """A tight label/value list. Cheaper vertically than real Word bullets."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.autofit = False
    for r, (label, value, color) in enumerate(rows):
        table.cell(r, 0).width = Inches(1.55)
        table.cell(r, 1).width = Inches(5.35)
        _text(table.cell(r, 0), label, size=9, bold=True, color=color)
        _text(table.cell(r, 1), value, size=9)


def _hours_verdict(doc, strongest, weakest, busy_losing) -> None:
    if not strongest and not weakest:
        _text(doc.add_paragraph(), "Too few decided offers in any hour to rank.", size=9, color=MUTED)
        return

    def phrase(hours: list[HourRow]) -> str:
        # "won X of Y decided", never "on Y offers" -- the rate denominator is
        # decided offers, and quoting the offer count beside a percentage the
        # offer count does not produce is how a reader concludes the maths is wrong.
        return ", ".join(f"{h.label} — {_pct(h.rate)} ({h.wins} of {h.decided})" for h in hours) or "—"

    rows = [
        ("Strongest", phrase(strongest), WIN_C),
        ("Weakest", phrase(weakest), LOSS_C),
    ]
    if busy_losing:
        rows.append((
            "Costing the most",
            ", ".join(f"{h.label} — {h.losses} lost of {h.offers}" for h in busy_losing),
            LOSS_C,
        ))
    _bullets(doc, rows)
    _text(
        doc.add_paragraph(),
        f"Hours with fewer than {MIN_HOURS_VOLUME} decided offers are excluded — "
        "on one day's volume they swing to 0% or 100% on a single job.",
        size=8, color=MUTED,
    )


def _trend_table(doc, stats: DayStats, baseline: list[DayStats]) -> None:
    """This day against the same weekday, most recent first."""
    if not baseline:
        _text(doc.add_paragraph(), "No earlier data for this weekday yet.", size=9, color=MUTED)
        return

    # header + the reported day + one row per baseline day + the average.
    # Allocating one fewer silently lands the average on top of the oldest
    # baseline row, which then reads as a 4-day average over 3 visible days.
    row_count = len(baseline) + 3
    table = doc.add_table(rows=row_count, cols=6)
    table.autofit = False
    widths = [1.55, 0.85, 0.85, 0.85, 0.95, 1.85]
    for c, w in enumerate(widths):
        for r in range(row_count):
            table.cell(r, c).width = Inches(w)
    for c, head in enumerate(["Date", "Offers", "Wins", "Losses", "Win rate", ""]):
        _text(table.cell(0, c), head, size=8, bold=True, color=MUTED, align="r" if c in (1, 2, 3, 4) else None)
        _border(table.cell(0, c), "bottom", "1C1C1C")

    series = [(stats, True)] + [(b, False) for b in baseline]
    for r, (day_stats, is_today) in enumerate(series, start=1):
        cells = [
            _short_date(day_stats.day),
            str(day_stats.offers), str(day_stats.wins), str(day_stats.losses), _pct(day_stats.rate),
            "◀ this report" if is_today else "",
        ]
        for c, value in enumerate(cells):
            _text(table.cell(r, c), value, size=9, bold=is_today,
                  align="r" if c in (1, 2, 3, 4) else None,
                  color=INK if is_today else MUTED)
            _border(table.cell(r, c), "top", RULE)
        if is_today:
            for c in range(6):
                _shade(table.cell(r, c), "F2F5F8")

    # Sits one past the last baseline row -- see the row_count note above.
    r = len(baseline) + 2
    avg = lambda pick: _avg([float(pick(b)) for b in baseline])  # noqa: E731
    summary = [
        f"Average of prior {len(baseline)}",
        f"{avg(lambda b: b.offers):.0f}",
        f"{avg(lambda b: b.wins):.0f}",
        f"{avg(lambda b: b.losses):.0f}",
        _pct(_avg([b.rate for b in baseline])),
        "",
    ]
    for c, value in enumerate(summary):
        _text(table.cell(r, c), value, size=8, bold=c == 4, color=MUTED,
              align="r" if c in (1, 2, 3, 4) else None)
        _border(table.cell(r, c), "top", "1C1C1C")


def _cancel_table(doc, stats: DayStats) -> None:
    """Cancellations by cause, most actionable first.

    Printed in CANCEL_KINDS order rather than by size, so the causes we can do
    something about lead even on a day when the generic bucket is the biggest.
    """
    order = [label for _raw, label, _blame in CANCEL_KINDS] + [CANCEL_OTHER]
    blame = {label: b for _raw, label, b in CANCEL_KINDS}
    rows = [(label, stats.cancel_kinds[label]) for label in order if stats.cancel_kinds[label]]
    if not rows:
        return

    table = doc.add_table(rows=len(rows) + 1, cols=4)
    table.autofit = False
    widths = [2.85, 0.70, 1.05, 1.70]
    for c, w in enumerate(widths):
        for r in range(len(rows) + 1):
            table.cell(r, c).width = Inches(w)
    for c, head in enumerate(["Cause", "Jobs", "Had a driver", ""]):
        _text(table.cell(0, c), head, size=8, bold=True, color=MUTED,
              align="r" if c in (1, 2) else None)
        _border(table.cell(0, c), "bottom", "1C1C1C")

    for r, (label, count) in enumerate(rows, start=1):
        ours = blame.get(label, False)
        engaged = stats.cancel_engaged[label]
        note = "within our control" if ours else ""
        if label == CANCEL_OTHER:
            note = "unrecognised label — check Towbook"
        for c, value in enumerate([label, str(count), str(engaged) if engaged else "—", note]):
            _text(table.cell(r, c), value, size=9, align="r" if c in (1, 2) else None,
                  bold=ours and c == 0, color=LOSS_C if (ours and c in (0, 3)) else INK)
            _border(table.cell(r, c), "top", RULE)

    if stats.cancel_reasons:
        _text(doc.add_paragraph(),
              "Our stated reason where one was recorded: "
              + ", ".join(f"{r} ({n})" for r, n in stats.cancel_reasons.most_common(3)),
              size=8, color=MUTED)


def _client_table(doc, stats: DayStats, limit: int = 6) -> None:
    ranked = sorted(stats.clients.items(), key=lambda kv: -kv[1].offers)[:limit]
    if not ranked:
        return
    # Cancelled gets its own column even though it is already inside Losses.
    # A client that cancels most of what it sends looks identical to a reliable
    # one in a Wins/Losses pair, and that is the single most useful thing this
    # table can tell the owner.
    table = doc.add_table(rows=len(ranked) + 1, cols=6)
    table.autofit = False
    widths = [2.30, 0.75, 0.70, 0.75, 0.95, 1.45]
    for c, w in enumerate(widths):
        for r in range(len(ranked) + 1):
            table.cell(r, c).width = Inches(w)
    for c, head in enumerate(["Client", "Offers", "Wins", "Losses", "of which cancelled", "Win rate"]):
        _text(table.cell(0, c), head, size=8, bold=True, color=MUTED, align="r" if c else None)
        _border(table.cell(0, c), "bottom", "1C1C1C")
    for r, (name, row) in enumerate(ranked, start=1):
        cold = row.rate is not None and row.rate < 0.25 and row.decided >= MIN_HOURS_VOLUME
        heavy = row.offers >= MIN_HOURS_VOLUME and row.canceled >= max(2, row.offers * 0.4)
        cells = [name, str(row.offers), str(row.wins), str(row.losses),
                 str(row.canceled) if row.canceled else "—", _pct(row.rate)]
        for c, value in enumerate(cells):
            _text(table.cell(r, c), value, size=9, align="r" if c else None,
                  bold=(heavy and c == 4) or (cold and c == 5),
                  color=LOSS_C if ((cold and c == 5) or (heavy and c == 4)) else INK)
            _border(table.cell(r, c), "top", RULE)


# -- document assembly -----------------------------------------------------


def _long_date(day: date) -> str:
    """"Sunday, August 9, 2026" without %-d, which is not portable to Windows."""
    return f"{day.strftime('%A, %B')} {day.day}, {day.year}"


def _short_date(day: date) -> str:
    return f"{day.strftime('%a %b')} {day.day}"


def _freshness_note(stats: DayStats, generated: datetime, tz: ZoneInfo) -> str | None:
    """Warn when the feed did not cover the whole reported day.

    A report that silently describes half a day is worse than no report: the
    win rate looks real and the hourly picture is simply missing its tail.
    """
    if not stats.has_data:
        return (
            "NO DATA for this day. Either no offers came in, or the overnight data pull "
            "did not run. Check before treating this as a quiet day."
        )
    last_local = stats.last_offer_utc.replace(tzinfo=timezone.utc).astimezone(tz)
    if last_local.date() < stats.day:
        return "Warning: the feed holds no offers for this day at all."
    if last_local.hour < 20:
        return (
            f"Warning: the last offer on record is {last_local.strftime('%H:%M')}. "
            "The day may be incompletely pulled, so evening hours can read as quiet "
            "when they were simply not collected."
        )
    return None


def build_document(
    stats: DayStats,
    prior: DayStats,
    baseline: list[DayStats],
    *,
    generated: datetime,
    tz: ZoneInfo,
):
    """Render the whole report. Three pages is a hard budget, not a target."""
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Segoe UI"
    normal.font.size = Pt(9)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.space_before = Pt(0)
    # Segoe UI must also be set for the East-Asian slot or Word substitutes a
    # fallback face for the block-bar glyphs in the hourly chart.
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "Segoe UI")

    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    strongest, weakest, busy_losing = _hour_findings(stats)

    # -- masthead ----------------------------------------------------------
    head = doc.add_paragraph()
    _text(head, "MORNING REPORT", size=8, bold=True, color=MUTED, space_after=1)
    title = doc.add_paragraph()
    _text(title, _long_date(stats.day), size=20, bold=True, space_after=1)
    sub = doc.add_paragraph()
    _text(
        sub,
        f"Business of the day prior · generated {generated.strftime('%a %b %d, %Y at %H:%M')} "
        f"· all times {tz.key}",
        size=8, color=MUTED, space_after=8,
    )

    warning = _freshness_note(stats, generated, tz)
    if warning:
        box = doc.add_table(rows=1, cols=1)
        cell = box.cell(0, 0)
        cell.width = Inches(7.3)
        _shade(cell, "FDECEA")
        _text(cell, warning, size=9, bold=True, color=LOSS_C)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    if not stats.has_data:
        return doc

    _scoreboard(doc, stats, prior, baseline)

    # -- the sentence that carries the report ------------------------------
    verdict = doc.add_table(rows=1, cols=1)
    vcell = verdict.cell(0, 0)
    vcell.width = Inches(7.3)
    _shade(vcell, "F2F5F8")
    _border(vcell, "left", "2F6FEB", size=18)
    _text(vcell, _headline(stats, baseline, busy_losing), size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # -- losses ------------------------------------------------------------
    _section(doc, "Where the losses went",
             f"{stats.losses} lost of {stats.decided} settled offers. Cancellations count as losses — "
             "a cancelled job is a job we did not run.")
    _bullets(doc, [
        ("Nobody answered", f"{stats.no_answer} offers expired with no response.", LOSS_C),
        ("We declined", f"{stats.declined} offers turned down." + (
            "  Top reasons: " + ", ".join(f"{r} ({n})" for r, n in stats.decline_reasons.most_common(3))
            if stats.decline_reasons else ""), LOSS_C),
        ("Cancelled", f"{stats.canceled} cancelled after the offer — broken out below.", LOSS_C),
    ])
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    if stats.canceled:
        _section(doc, "Why the cancellations happened",
                 "“Had a driver” means the job was already crewed when it died — a broken commitment "
                 "or a wasted roll, not a client changing its mind.")
        _cancel_table(doc, stats)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # -- hourly ------------------------------------------------------------
    _section(doc, "How the day flowed, hour by hour",
             "Local time. Bars show offer volume. Rate is Won ÷ (Won + Lost), and Lost includes "
             "cancellations. Won and Lost fall short of Off only where an offer is still pending.")
    _hourly_table(doc, stats)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    _section(doc, "Strongest and weakest hours")
    _hours_verdict(doc, strongest, weakest, busy_losing)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # -- trend -------------------------------------------------------------
    _section(doc, f"Trend — this {stats.day.strftime('%A')} against recent {stats.day.strftime('%A')}s",
             "Same weekday only. Tow volume runs on a weekly cycle, so comparing across weekdays "
             "invents swings that are only the calendar.")
    _trend_table(doc, stats, baseline)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    _section(doc, "By client", "Ranked by offers received. Counts are jobs, not dollars.")
    _client_table(doc, stats)

    foot = doc.add_paragraph()
    _text(foot, "Counts are job counts. This account reports no offer amounts, so no revenue figure "
                "is shown or implied.", size=7, color=MUTED)
    return doc


# -- output ----------------------------------------------------------------


def month_folder(root: Path, day: date) -> Path:
    """``<root>/2026-08 August`` -- sorts chronologically, reads like a month."""
    return root / f"{day.strftime('%Y-%m')} {day.strftime('%B')}"


def report_path(root: Path, day: date) -> Path:
    """Named for the day it DESCRIBES, not the day it ran.

    The 6 AM run on the 10th reports the 9th and is filed as the 9th; filing it
    under the run date would put every report one folder off at a month boundary.
    """
    return month_folder(root, day) / f"{day.strftime('%Y-%m-%d %A')}.docx"


def generate(
    db_path: Path,
    out_root: Path,
    *,
    day: date | None = None,
    tz_name: str = "America/Detroit",
    now: datetime | None = None,
) -> Path:
    """Build one report and return the path written."""
    tz = ZoneInfo(tz_name)
    generated = now or datetime.now(tz)
    target = day or (generated.astimezone(tz).date() - timedelta(days=1))

    conn = sqlite3.connect(f"file:{db_path}?mode=ro", uri=True)
    try:
        stats = load_day(conn, target, tz)
        prior = load_day(conn, target - timedelta(days=1), tz)
        baseline, _ = same_weekday_baseline(conn, target, tz)
    finally:
        conn.close()

    doc = build_document(stats, prior, baseline, generated=generated, tz=tz)
    path = report_path(out_root, target)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path
