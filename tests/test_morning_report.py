"""Tests for the printable 6 AM report.

These cover the four things that would be WRONG BUT PLAUSIBLE on the page --
the failure mode that matters here, because a report that prints cleanly and
lies is worse than one that crashes. Nobody re-derives the numbers by hand.

    1. Local-hour conversion. Stored hours are UTC; a report that shows them
       raw is off by the whole offset and names the wrong shift.
    2. Day boundaries. "Yesterday" is local midnight to local midnight.
    3. Client cancellations are not losses and are out of the rate denominator.
    4. The trend table allocates a row per baseline day PLUS the average.
"""

from __future__ import annotations

import sqlite3
from datetime import date, datetime, timedelta, timezone
from zoneinfo import ZoneInfo

import pytest

from towbook_agent.agents import morning_report as mr

TZ = ZoneInfo("America/Detroit")


@pytest.fixture()
def conn() -> sqlite3.Connection:
    """A stand-in requests table holding only the columns the report reads."""
    db = sqlite3.connect(":memory:")
    db.execute(
        """
        CREATE TABLE requests (
            offered_at TEXT, status TEXT, denial_reason TEXT, client_name TEXT,
            status_raw TEXT, driver_assigned TEXT
        )
        """
    )
    return db


def add(
    db: sqlite3.Connection,
    local: datetime,
    status: str,
    *,
    client: str = "Agero",
    reason=None,
    raw=None,
    driver=None,
) -> None:
    """Insert one offer given its LOCAL time, stored as naive UTC like ingestion does."""
    utc = local.replace(tzinfo=TZ).astimezone(timezone.utc).replace(tzinfo=None)
    db.execute(
        "INSERT INTO requests VALUES (?,?,?,?,?,?)",
        (utc.isoformat(sep=" "), status, reason, client, raw, driver),
    )


# -- 1 & 2: time handling --------------------------------------------------


def test_offers_land_in_their_local_hour_not_the_stored_utc_hour(conn):
    """09:00 local must appear at 09:00 even though it is stored as 13:00 UTC."""
    add(conn, datetime(2026, 7, 15, 9, 30), "accepted")
    stats = mr.load_day(conn, date(2026, 7, 15), TZ)

    assert stats.hours[9].offers == 1
    assert stats.hours[13].offers == 0, "read the stored UTC hour instead of the local one"


def test_day_window_is_local_midnight_to_local_midnight(conn):
    """The last minute of the local day belongs to it; the next minute does not."""
    add(conn, datetime(2026, 7, 15, 0, 0), "accepted")
    add(conn, datetime(2026, 7, 15, 23, 59), "accepted")
    add(conn, datetime(2026, 7, 14, 23, 59), "accepted")
    add(conn, datetime(2026, 7, 16, 0, 1), "accepted")

    assert mr.load_day(conn, date(2026, 7, 15), TZ).offers == 2


def test_local_day_bounds_absorb_a_dst_transition():
    """Spring forward is a 23-hour local day. The window must still be exact."""
    start, end = mr._local_day_bounds(date(2026, 3, 8), TZ)  # US spring-forward
    assert (end - start) == timedelta(hours=23)


# -- 3: outcome vocabulary -------------------------------------------------


def test_cancellations_count_as_losses(conn):
    """3 wins, 1 expiry, 6 cancels is 30% -- a cancelled job is one we did not run.

    The old rule excluded cancellations and reported this same day as 75%. On
    the real record that flattery was 12 points overall and 39 points for
    Allstate, whose offers cancel more often than not.
    """
    for _ in range(3):
        add(conn, datetime(2026, 7, 15, 10, 0), "accepted")
    add(conn, datetime(2026, 7, 15, 10, 0), "expired")
    for _ in range(6):
        add(conn, datetime(2026, 7, 15, 10, 0), "canceled", raw="Cancelled")

    stats = mr.load_day(conn, date(2026, 7, 15), TZ)
    assert (stats.wins, stats.losses, stats.canceled) == (3, 7, 6)
    assert stats.decided == 10
    assert stats.rate == pytest.approx(0.30)


def test_cancellations_are_attributed_to_a_cause(conn):
    """"Why" is the point -- a bare total would just move the problem."""
    add(conn, datetime(2026, 7, 15, 10, 0), "canceled", raw="Service Failure Confirmed", driver="Unit 7")
    add(conn, datetime(2026, 7, 15, 10, 0), "canceled", raw="Rejected By Motor Club", reason="No Drivers Available")
    add(conn, datetime(2026, 7, 15, 10, 0), "canceled", raw="Goa Approved By Motor Club", driver="Unit 3")

    stats = mr.load_day(conn, date(2026, 7, 15), TZ)
    assert stats.cancel_kinds["We failed after accepting"] == 1
    assert stats.cancel_kinds["Club pulled it back"] == 1
    assert stats.cancel_kinds["Gone on arrival"] == 1
    assert stats.cancel_engaged["We failed after accepting"] == 1
    assert stats.cancel_engaged["Club pulled it back"] == 0
    assert stats.cancel_reasons["No Drivers Available"] == 1


def test_an_unrecognised_cancellation_label_is_surfaced_not_dropped(conn):
    """If Towbook adds a label it must appear as unclassified, never vanish."""
    add(conn, datetime(2026, 7, 15, 10, 0), "canceled", raw="Some New Towbook Status")

    stats = mr.load_day(conn, date(2026, 7, 15), TZ)
    assert stats.cancel_kinds[mr.CANCEL_OTHER] == 1
    assert sum(stats.cancel_kinds.values()) == stats.canceled


def test_losses_split_into_unanswered_and_declined(conn):
    add(conn, datetime(2026, 7, 15, 10, 0), "expired")
    add(conn, datetime(2026, 7, 15, 10, 0), "denied", reason="Equipment Not Available")

    stats = mr.load_day(conn, date(2026, 7, 15), TZ)
    assert (stats.no_answer, stats.declined) == (1, 1)
    assert stats.decline_reasons["Equipment Not Available"] == 1


def test_a_day_with_no_offers_has_no_rate_rather_than_zero(conn):
    """0% would claim we turned down work that was never offered."""
    stats = mr.load_day(conn, date(2026, 7, 15), TZ)
    assert stats.rate is None
    assert not stats.has_data


# -- 4: presentation invariants -------------------------------------------


def test_rising_losses_are_coloured_bad_even_though_the_arrow_points_up():
    up_good, colour_good = mr._delta(10, 5)
    up_bad, colour_bad = mr._delta(10, 5, higher_is_better=False)

    assert up_good.startswith("▲") and up_bad.startswith("▲")
    assert colour_good == mr.WIN_C
    assert colour_bad == mr.LOSS_C, "more losses must never print green"


def test_hour_superlatives_ignore_thin_hours(conn):
    """A 1-of-1 hour is 100% and must not outrank a real one."""
    add(conn, datetime(2026, 7, 15, 3, 0), "accepted")  # 1 decided -> excluded
    for _ in range(8):
        add(conn, datetime(2026, 7, 15, 10, 0), "accepted")
    add(conn, datetime(2026, 7, 15, 10, 0), "expired")

    strongest, _, _ = mr._hour_findings(mr.load_day(conn, date(2026, 7, 15), TZ))
    assert [h.hour for h in strongest] == [10]


def test_trend_table_keeps_every_baseline_day_and_the_average(conn):
    """Under-allocating rows lands the average on top of the oldest day."""
    from docx import Document

    for offset in range(0, 36, 7):  # the day itself plus 4 prior same-weekdays
        add(conn, datetime(2026, 7, 31, 10, 0) - timedelta(days=offset), "accepted")

    day = date(2026, 7, 31)
    stats = mr.load_day(conn, day, TZ)
    baseline, _ = mr.same_weekday_baseline(conn, day, TZ)
    assert len(baseline) == 4

    doc = Document()
    mr._trend_table(doc, stats, baseline)
    table = doc.tables[0]

    assert len(table.rows) == len(baseline) + 3  # header + today + 4 + average
    assert table.cell(len(baseline) + 2, 0).text.startswith("Average of prior")
    dates = [table.cell(r, 0).text for r in range(1, len(baseline) + 2)]
    assert len(set(dates)) == 5, "a baseline row was overwritten"


def test_baseline_skips_days_with_no_data_rather_than_averaging_in_zeros(conn):
    """A gap in the feed must not drag the baseline toward nothing."""
    add(conn, datetime(2026, 7, 31, 10, 0), "accepted")
    add(conn, datetime(2026, 7, 24, 10, 0), "accepted")  # 17th and 10th missing
    add(conn, datetime(2026, 7, 3, 10, 0), "accepted")

    baseline, _ = mr.same_weekday_baseline(conn, date(2026, 7, 31), TZ)
    assert [b.day for b in baseline] == [date(2026, 7, 24), date(2026, 7, 3)]


# -- output shape ----------------------------------------------------------


def test_report_is_filed_by_the_day_it_describes(tmp_path):
    path = mr.report_path(tmp_path, date(2026, 8, 9))
    assert path.parent.name == "2026-08 August"
    assert path.name == "2026-08-09 Sunday.docx"


def test_generate_writes_a_readable_docx_even_for_an_empty_day(tmp_path):
    """The 6 AM job must produce a file when the pull failed, not nothing."""
    from docx import Document

    db = tmp_path / "t.db"
    con = sqlite3.connect(db)
    con.execute(
        "CREATE TABLE requests (offered_at TEXT, status TEXT, denial_reason TEXT,"
        " client_name TEXT, status_raw TEXT, driver_assigned TEXT)"
    )
    con.commit()
    con.close()

    path = mr.generate(db, tmp_path / "out", day=date(2026, 8, 5))
    assert path.exists()
    text = "\n".join(p.text for p in Document(str(path)).paragraphs)
    assert "NO DATA" in "\n".join(
        c.text for t in Document(str(path)).tables for r in t.rows for c in r.cells
    ) or "NO DATA" in text
